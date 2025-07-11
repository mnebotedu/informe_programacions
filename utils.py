from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def afegir_titol(doc, text):
    titol = doc.add_heading(text, 0)
    titol.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def afegir_encap1(doc, text):
    doc.add_heading(text, level=1)


def afegir_text(doc, text):
    doc.add_paragraph(text)


def afegir_llista(doc, elements):
    for el in elements:
        doc.add_paragraph(f"{el}", style='List Bullet')


def afegir_seccio_buida(doc, titol):
    doc.add_heading(titol, level=1)
    taula = doc.add_table(rows=1, cols=1)
    taula.style = 'Table Grid'
    taula.cell(0, 0).text = "\n\n\n"


def afegir_seccio_amb_indicacions(doc, titol, indicacions):
    doc.add_heading(titol, level=1)
    taula = doc.add_table(rows=2, cols=1)
    taula.style = 'Table Grid'

    # Cel·la d'indicacions
    cell1 = taula.cell(0, 0)
    p = cell1.paragraphs[0]
    run = p.add_run(indicacions)
    run.italic = True
    font = run.font
    font.color.rgb = RGBColor(80, 80, 80)  # gris suau

    # Cel·la editable
    taula.cell(1, 0).text = "\n\n\n"


def afegir_seccio_amb_text_i_buit(doc, titol, text):
    doc.add_heading(titol, level=1)

    # Text introductori suau
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = True
    font = run.font
    font.color.rgb = RGBColor(100, 100, 100)

    # Taula 1x1 sense vores
    taula = doc.add_table(rows=1, cols=1)
    cell = taula.cell(0, 0)
    cell.text = "\n\n\n"

    # Eliminar les vores
    tbl = taula._element
    for border_dir in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        borders = tbl.xpath(f'.//w:{border_dir}')
        for b in borders:
            b.getparent().remove(b)



def generar_docx(dades, nom_fitxer='informe.docx'):
    doc = Document()

    # Afegim una taula a l'inici amb els dos logos (simula una capçalera visual)
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import os

    taula_logos = doc.add_table(rows=1, cols=2)
    taula_logos.autofit = True

    cel_esq = taula_logos.cell(0, 0)
    cel_dre = taula_logos.cell(0, 1)

    try:
        cel_esq.paragraphs[0].add_run().add_picture(
            os.path.join('static', 'logo_conselleria_color.png'), width=Inches(1.5))
    except:
        cel_esq.text = "[Logo conselleria]"

    try:
        par_dre = cel_dre.paragraphs[0]
        par_dre.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        par_dre.add_run().add_picture(
            os.path.join('static', 'logo_centre.png'), width=Inches(1.5))
    except:
        cel_dre.text = "[Logo centre educatiu]"


    # Títol principal
    afegir_titol(doc, 'Esborrany de programació didàctica')


    # Informació bàsica
    info_basica = [
        f"Nivell educatiu: {dades['nivell']}",
        f"Curs: {dades['curs']}",
        f"Matèria: {dades['materia']}"
    ]
    afegir_llista(doc, info_basica)

    # Secció 1
    afegir_encap1(doc, '1. LES COMPETÈNCIES ESPECÍFIQUES I ELS CRITERIS D’AVALUACIÓ DE LA MATÈRIA O ÀMBIT PER A TOT EL CURS.')
    afegir_text(doc, "Competències específiques seleccionades:")
    afegir_llista(doc, dades['competencies'])

    afegir_text(doc, "Criteris d'avaluació seleccionats:")
    afegir_llista(doc, dades['criteris'])

    # Secció 2
    afegir_encap1(doc, '2. ELS SABERS BÀSICS SEQÜENCIATS I TEMPORITZATS PERA A CADA CURS, ORGANITZATS EN UNITATS DE PROGRAMACIÓ.')
    afegir_text(doc, "Sabers bàsics seleccionats:")
    afegir_llista(doc, dades['sabers'])
    if dades.get('concrecions'):
        afegir_text(doc, "Concrecions dels sabers seleccionats:")
        afegir_llista(doc, dades['concrecions'])



    # Seccions 3 a 10 amb text introductori
    seccions_textos = [
        ("3. PROCESSOS I INSTRUMENTS D’AVALUACIÓ I SISTEMA DE QUALIFICACIÓ.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de descriure com organitzaràs l’avaluació del teu alumnat, assegurant que sigui global, contínua, formativa, "
         "individualitzada i orientadora. Detalla com es desenvoluparà al llarg del curs: identifica els moments d’avaluació (inicial, formativa i "
         "sumativa) i les finalitats que persegueixes (conèixer el progrés dels infants, adaptar la teva metodologia, detectar necessitats de suport i " 
         "millorar la pràctica docent). Explica també què avaluaràs (aprenentatges de l’alumnat, la teva tasca com a docent i la mateixa programació) i " 
         "quins agents participaran (alumnat, equip docent, famílies…).\n"
         "Indica clarament els instruments que faràs servir per recollir evidències: rúbriques (per a autoavaluació, coavaluació i heteroavaluació), llistes " 
         "de control, portafolis, observació directa o enquestes de valoració. Finalment, especifica el sistema de qualificació qualitativa que aplicaràs, "
         "detallant com definiràs els nivells d’assoliment dels criteris d’avaluació i quina ponderació assignaràs a cada un d’ells, seguint la normativa vigent."),
        ("4. ESTRATÈGIES DIDÀCTIQUES I METODOLÒGIQUES: ORGANITZACIÓ, RECURSOS, AGRUPAMENTS, CRITERIS PER A L’ELABORACIÓ DE SITUACIONS I ACTIVITATS QUE ES CONSIDERIN NECESSÀRIES.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de descriure les opcions metodològiques que guiaran la teva pràctica docent, assegurant que estiguin alineades amb els principis pedagògics de l’etapa "
         "i amb l’enfocament competencial del currículum. Detalla els fonaments i principis metodològics que aplicaràs (com ara l’aprenentatge actiu, cooperatiu, significatiu o per "
         "descobriment), i explica com s’organitzaran el temps, l’espai i els agrupaments (individuals, parelles, petits grups, gran grup) per afavorir l’aprenentatge de tot l’alumnat.\n" 
         "Inclou els recursos humans i materials que preveus utilitzar —des de materials manipulatius i digitals fins a espais externs— i concreta els tipus d’activitats d’ensenyament-aprenentatge "
         "que desplegaràs: activitats d’iniciació i presa de consciència, de desenvolupament i aprofundiment, de retroacció i de síntesi o transferència. Finalment, explica els criteris que faràs "
         "servir per dissenyar activitats i tasques contextualitzades, rellevants i significatives per als infants."),
        ("5. ACTUACIONS GENERALS D’ATENCIÓ A LES NECESSITATS INDIVIDUALS.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de concretar les mesures que aplicaràs per atendre la diversitat d’estils, ritmes i necessitats d’aprenentatge del teu alumnat, amb l’objectiu de garantir la "
         "seva inclusió i participació. Detalla les mesures universals que aplicaràs a tot el grup classe, basades en els principis del Disseny Universal per a l’Aprenentatge (DUA), tenint "
         "en compte múltiples formes de representació de la informació, d’expressió de l’aprenentatge i d’implicació en les tasques.\n A més, explica quines mesures addicionals o intensives "
         "preveus per als casos en què les universals no siguin suficients. Aquestes poden incloure adaptacions metodològiques, suports específics o intervencions personalitzades. Finalment, "
         "recorda descriure com es durà a terme la coordinació amb altres professionals del centre (tutors, PT, AL, EOEP...) per garantir una resposta educativa coherent, inclusiva i ajustada a cada realitat."),
        ("6. CONCRECIÓ DELS ELEMENTS TRANSVERSALS ESTABLERTS EN EL PROJECTE EDUCATIU.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de descriure com integraràs els elements transversals establerts en el Projecte Educatiu del Centre (PEC) dins la teva matèria. Explica de manera concreta com "
         "treballaràs valors com la coeducació, la igualtat d’oportunitats, l’educació per a la ciutadania, la convivència democràtica, la sostenibilitat ambiental, la promoció de la salut i el "
         "respecte per la diversitat. És important que indiquis com aquests valors s’incorporaran a les pràctiques didàctiques i a les situacions d’aprenentatge, promovent actituds crítiques, "
         "responsables i compromeses amb l’entorn i amb la societat."),
        ("7. ACTIVITATS COMPLEMENTÀRIES.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has d’indicar les activitats complementàries que tens previstes al llarg del curs, relacionades amb la programació didàctica. Es tracta d’activitats que, tot i no ser "
         "estrictament curriculars, contribueixen de manera significativa al desenvolupament de les competències clau i a l’enriquiment de l’experiència educativa. Pots incloure sortides "
         "escolars, visites culturals, tallers externs, activitats lúdiques amb finalitats didàctiques o participació en esdeveniments comunitaris. Assegura’t que aquestes propostes siguin "
         "coherents amb els objectius de l’àrea i adaptades a les característiques i necessitats del teu alumnat."),
        ("8. PLA DE SEGUIMENT ALS ALUMNES QUE NO PROMOCIONEN.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de descriure les accions, mesures i suports específics que preveus per a l’alumnat que no promociona de curs. Aquest pla de seguiment ha d’incloure intervencions personalitzades, "
         "adaptades a les seves necessitats educatives concretes, amb l’objectiu de facilitar-ne la reincorporació progressiva al grup i garantir l’assoliment de les competències pendents. Pots incloure mesures "
         "com tutories individualitzades, reforç educatiu, adaptacions curriculars, orientació específica o coordinació amb altres professionals del centre per fer un seguiment global i efectiu."),
        ("9. MECANISME DE REVISIÓ, AVALUACIÓ I MODIFICACIÓ DE LES PROGRAMACIONS DIDÀCTIQUES.",
         "Per emplenar aquest apartat, posa atenció en les indicacions que et proposam. Pots esborrar aquesta part en acabar:\n\n"
         "En aquest apartat has de descriure el procediment que seguiràs per revisar, avaluar i actualitzar la programació didàctica al final del curs. Explica com analitzaràs la pertinència i eficàcia de les "
         "estratègies metodològiques, activitats i criteris d’avaluació, així com l’adequació dels sabers bàsics a la realitat de l’alumnat. Aquest procés ha de permetre identificar àrees de millora, detectar "
         "necessitats d’ajust i proposar canvis metodològics o organitzatius. També cal que especifiquis com es farà la coordinació amb l’equip docent per acordar modificacions i garantir que la nova versió de la "
         "programació sigui coherent amb la normativa vigent, el projecte educatiu del centre i les necessitats reals de l’aula."),

    ]

    for titol, text in seccions_textos:
        afegir_seccio_amb_text_i_buit(doc, titol, text)

    doc.save(nom_fitxer)
